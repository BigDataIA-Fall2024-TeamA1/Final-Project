from fastapi import FastAPI, HTTPException
from dotenv import load_dotenv
from pydantic import BaseModel
import openai
import os
import requests
from docx import Document  # For Word document generation
from fpdf import FPDF  # For PDF generation

# Load environment variables
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
SERPAPI_KEY = os.getenv("SERPAPI_KEY")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME")

# Initialize OpenAI and Pinecone
openai.api_key = OPENAI_API_KEY

# FastAPI Initialization
app = FastAPI()

class SearchRequest(BaseModel):
    query: str

class CaseQueryRequest(BaseModel):
    query: str

class ReportMetadata(BaseModel):
    lawyer_name: str
    law_firm: str
    court: str
    jurisdiction: str
    facts: str
    issues: str
    reasoning: str
    decision: str
    feedback: str = None  # Optional field for user feedback

# Report Generation Agent
class ReportGenerationAgent:
    """Generate and handle professional legal reports."""
    
    def generate_report(self, metadata: ReportMetadata):
        """Generate a professional legal report using OpenAI."""
        feedback_prompt = f" Feedback from user: {metadata.feedback}" if metadata.feedback else ""
        prompt = f"""
        Generate a professional legal report based on the following metadata:

        Lawyer Name: {metadata.lawyer_name}
        Law Firm: {metadata.law_firm}
        Court: {metadata.court}
        Jurisdiction: {metadata.jurisdiction}
        Facts: {metadata.facts}
        Issues: {metadata.issues}
        Reasoning: {metadata.reasoning}
        Decision: {metadata.decision}
        {feedback_prompt}

        Structure the report as follows:
        1. Introduction: Explain the purpose of the report.
        2. Statement of Facts: Detail the relevant facts of the case.
        3. Discussion: Provide a detailed analysis and legal reasoning.
        4. Conclusion: Summarize findings and recommendations.
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a professional legal document generator."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.5
        )
        return response['choices'][0]['message']['content']

    def save_to_word(self, content: str, file_name: str = "legal_report.docx"):
        """Save report content to a Word document."""
        doc = Document()
        doc.add_heading("Legal Report", level=1)
        doc.add_paragraph(content)
        file_path = f"./{file_name}"
        doc.save(file_path)
        return file_path

    def save_to_pdf(self, content: str, file_name: str = "legal_report.pdf"):
        """Save report content to a PDF file."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, content)
        file_path = f"./{file_name}"
        pdf.output(file_path)
        return file_path

class WebSearchAgent:
    """Search legal resources using SERPAPI."""
    def search(self, query: str):
        serpapi_url = "https://serpapi.com/search"
        params = {
            "q": query,
            "api_key": SERPAPI_KEY,
            "engine": "google",
            "site": "thomsonreuters.com,law360.com"  # Restrict to legal sites
        }
        response = requests.get(serpapi_url, params=params)
        if response.status_code != 200:
            raise ValueError("Failed to fetch search results from SERPAPI.")
        return response.json().get("organic_results", [])


class CaseRetrievalAgent:
    """Retrieve relevant court cases."""
    def retrieve_cases(self, query: str):
        court_listener_url = "https://www.courtlistener.com/api/rest/v3/search/"
        params = {"q": query}
        response = requests.get(court_listener_url, params=params)
        if response.status_code != 200:
            raise ValueError("Failed to fetch cases from Court Listener.")
        cases = response.json().get("results", [])
        return [
            {"title": case["caseName"], "url": f"https://www.courtlistener.com{case['absolute_url']}"}
            for case in cases
        ]


class StrategyGenerationAgent:
    """Generate a comprehensive legal strategy."""
    
    def generate_strategy(self, metadata: dict):
        """Use OpenAI to generate a detailed legal strategy."""
        prompt = f"""
        Based on the following case details, generate a detailed legal strategy:

        - **Facts**: {metadata.get('facts')}
        - **Issues**: {metadata.get('issues')}
        - **Reasoning**: {metadata.get('reasoning')}
        - **Decision**: {metadata.get('decision')}

        Provide a comprehensive strategy that includes:
        1. Recommended legal approaches to address the case.
        2. Alternative actions or strategies that could have been considered.
        3. Potential arguments from the opposing party and how to counter them.
        4. Advice on whether the case should proceed to trial, be settled, or require filing additional motions.
        5. Specific legal precedents or similar cases that support your strategy.

        Ensure the response is detailed and actionable, tailored for a professional legal audience.
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert legal strategist providing actionable advice."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.3  # Lower temperature for more deterministic responses
        )
        return response['choices'][0]['message']['content']



@app.post("/solution2/report")
async def solution2_report(metadata: ReportMetadata):
    """Generate and return a legal report."""
    try:
        agent = ReportGenerationAgent()
        report_content = agent.generate_report(metadata)
        word_file = agent.save_to_word(report_content)
        pdf_file = agent.save_to_pdf(report_content)
        return {
            "report_content": report_content,  # Return the content for frontend display
            "word_file": word_file,
            "pdf_file": pdf_file
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/solution2/search")
async def solution2_search(request: SearchRequest):
    """Search legal resources using SERPAPI."""
    try:
        agent = WebSearchAgent()
        results = agent.search(request.query)
        return {"results": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/solution2/cases")
async def solution2_cases(request: CaseQueryRequest):
    """Retrieve relevant court cases."""
    try:
        agent = CaseRetrievalAgent()
        cases = agent.retrieve_cases(request.query)  # Extract the `query` field
        return {"cases": cases}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/solution2/strategy")
async def solution2_strategy(metadata: dict):
    """Generate a legal strategy."""
    try:
        agent = StrategyGenerationAgent()
        strategy = agent.generate_strategy(metadata)
        return {"strategy": strategy}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@app.get("/solution2/download/{file_name}")
async def download_file(file_name: str):
    """Download generated report."""
    file_path = f"./{file_name}"
    if os.path.exists(file_path):
        return FileResponse(file_path, media_type="application/octet-stream", filename=file_name)
    raise HTTPException(status_code=404, detail="File not found.")
